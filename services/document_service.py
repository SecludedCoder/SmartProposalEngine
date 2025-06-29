#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
文件路径: smart_proposal_engine/services/document_service.py
功能说明: 文档处理服务，支持多种文档格式的读取和处理
作者: SmartProposal Team
创建日期: 2025-06-27
最后修改: 2025-06-27
版本: 1.0.0
"""

import os
import sys
import time
from typing import Dict, List, Optional, Union, Tuple
from datetime import datetime
from pathlib import Path
import mimetypes

# 文档处理库
try:
    import docx
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("警告：python-docx未安装，无法处理DOCX文件")

try:
    import PyPDF2
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("警告：PyPDF2未安装，无法处理PDF文件")

# 添加项目根目录到系统路径
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from services.base_service import BaseService, ProcessingResult
from utils.file_utils import get_file_extension, get_file_metadata, format_file_size
from utils.format_utils import clean_text, format_metadata_display


class DocumentService(BaseService):
    """
    文档处理服务
    
    主要功能:
    1. 支持DOCX文档解析（使用python-docx）
    2. 支持PDF文档解析（使用PyPDF2）
    3. 支持TXT文档读取
    4. 统一输出为标准文本格式
    5. 保留文档元数据（页数、创建时间等）
    
    使用示例:
        service = DocumentService()
        result = service.process(file_path, options={'extract_metadata': True})
    """
    
    # 支持的文档格式
    SUPPORTED_FORMATS = {
        '.docx': 'Microsoft Word Document',
        '.doc': 'Microsoft Word Document (Legacy)',
        '.pdf': 'PDF Document',
        '.txt': 'Text File',
        '.rtf': 'Rich Text Format',
        '.odt': 'OpenDocument Text'
    }
    
    def __init__(self):
        super().__init__()
        self.check_dependencies()
    
    def check_dependencies(self):
        """检查依赖库是否可用"""
        self.capabilities = {
            'docx': DOCX_AVAILABLE,
            'pdf': PDF_AVAILABLE,
            'txt': True  # 总是可用
        }
        
        if not DOCX_AVAILABLE:
            print("提示：安装python-docx以支持DOCX文件处理: pip install python-docx")
        if not PDF_AVAILABLE:
            print("提示：安装PyPDF2以支持PDF文件处理: pip install PyPDF2")
    
    def get_available_templates(self) -> List[str]:
        """获取可用的模板列表（文档服务不使用模板）"""
        return []
    
    def validate_input(self, input_data: Union[str, bytes, Path]) -> bool:
        """验证输入是否为支持的文档格式"""
        if isinstance(input_data, (str, Path)):
            ext = get_file_extension(input_data)
            return ext.lower() in self.SUPPORTED_FORMATS
        return False
    
    def process(self,
                input_data: Union[str, bytes, Path],
                template: Optional[str] = None,
                options: Optional[Dict] = None) -> ProcessingResult:
        """
        处理文档文件
        
        Args:
            input_data: 文档文件路径
            template: 未使用（保持接口一致）
            options: 处理选项
                - extract_metadata: 是否提取元数据
                - preserve_formatting: 是否保留格式信息
                - clean_output: 是否清理输出文本
                - progress_callback: 进度回调函数
        
        Returns:
            ProcessingResult: 处理结果
        """
        start_time = time.time()
        options = options or {}
        progress_callback = options.get('progress_callback')
        
        try:
            # 确保输入是文件路径
            file_path = Path(input_data)
            if not file_path.exists():
                raise FileNotFoundError(f"文件不存在: {file_path}")
            
            # 获取文件信息
            file_metadata = get_file_metadata(file_path)
            ext = file_metadata['extension'].lower()
            
            if progress_callback:
                progress_callback(f"正在处理 {file_metadata['filename']}...")
            
            # 根据文件类型选择处理方法
            if ext == '.docx':
                content, doc_metadata = self._process_docx(file_path, options)
            elif ext == '.pdf':
                content, doc_metadata = self._process_pdf(file_path, options)
            elif ext in ['.txt', '.text']:
                content, doc_metadata = self._process_text(file_path, options)
            elif ext == '.doc':
                # 旧版Word文档，尝试作为DOCX处理（可能需要转换）
                content, doc_metadata = self._process_legacy_doc(file_path, options)
            else:
                # 尝试作为纯文本处理
                content, doc_metadata = self._process_text(file_path, options)
            
            # 清理文本（如果需要）
            if options.get('clean_output', True):
                content = clean_text(content)
            
            # 构建元数据
            metadata = {
                **file_metadata,
                **doc_metadata,
                'processing_options': {
                    'extract_metadata': options.get('extract_metadata', True),
                    'preserve_formatting': options.get('preserve_formatting', False),
                    'clean_output': options.get('clean_output', True)
                }
            }
            
            # 添加内容统计
            metadata.update(self._analyze_content(content))
            
            if progress_callback:
                progress_callback(f"文档处理完成，共提取 {metadata.get('word_count', 0)} 个词")
            
            processing_time = time.time() - start_time
            
            return ProcessingResult(
                content=content,
                metadata=metadata,
                source_type='document',
                processing_time=processing_time,
                model_used='',  # 文档处理不使用AI模型
                tokens_consumed={}
            )
            
        except Exception as e:
            if progress_callback:
                progress_callback(f"文档处理失败: {str(e)}")
            
            processing_time = time.time() - start_time
            
            return ProcessingResult(
                content='',
                metadata={
                    'error': str(e),
                    'file_path': str(input_data)
                },
                source_type='document',
                processing_time=processing_time,
                model_used='',
                tokens_consumed={},
                error=str(e)
            )
    
    def _process_docx(self, file_path: Path, options: Dict) -> Tuple[str, Dict]:
        """处理DOCX文档"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx库未安装，无法处理DOCX文件")
        
        doc = Document(str(file_path))
        content_parts = []
        metadata = {
            'format': 'docx',
            'paragraph_count': 0,
            'table_count': len(doc.tables),
            'image_count': len(doc.inline_shapes)
        }
        
        # 提取核心属性（如果可用）
        if options.get('extract_metadata', True):
            try:
                core_props = doc.core_properties
                metadata.update({
                    'title': core_props.title or '',
                    'author': core_props.author or '',
                    'subject': core_props.subject or '',
                    'keywords': core_props.keywords or '',
                    'created': core_props.created.isoformat() if core_props.created else '',
                    'modified': core_props.modified.isoformat() if core_props.modified else '',
                    'last_modified_by': core_props.last_modified_by or ''
                })
            except:
                pass
        
        # 提取段落文本
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                content_parts.append(text)
                metadata['paragraph_count'] += 1
        
        # 提取表格文本（如果需要）
        if options.get('extract_tables', True) and doc.tables:
            content_parts.append("\n[表格内容]\n")
            for table_idx, table in enumerate(doc.tables):
                content_parts.append(f"\n表格 {table_idx + 1}:")
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        content_parts.append(row_text)
        
        content = "\n\n".join(content_parts)
        
        # 计算页数（估算）
        # DOCX没有固定的页数概念，这里根据字数估算
        word_count = len(content.split())
        estimated_pages = max(1, word_count // 250)  # 假设每页250词
        metadata['page_count'] = estimated_pages
        metadata['page_count_note'] = 'estimated'
        
        return content, metadata
    
    def _process_pdf(self, file_path: Path, options: Dict) -> Tuple[str, Dict]:
        """处理PDF文档"""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2库未安装，无法处理PDF文件")
        
        content_parts = []
        metadata = {
            'format': 'pdf',
            'page_count': 0,
            'encrypted': False
        }
        
        try:
            with open(file_path, 'rb') as pdf_file:
                pdf_reader = PdfReader(pdf_file)
                
                # 检查是否加密
                if pdf_reader.is_encrypted:
                    metadata['encrypted'] = True
                    # 尝试使用空密码解密
                    if not pdf_reader.decrypt(''):
                        raise ValueError("PDF文件已加密，需要密码")
                
                metadata['page_count'] = len(pdf_reader.pages)
                
                # 提取元数据
                if options.get('extract_metadata', True) and pdf_reader.metadata:
                    pdf_meta = pdf_reader.metadata
                    metadata.update({
                        'title': pdf_meta.get('/Title', ''),
                        'author': pdf_meta.get('/Author', ''),
                        'subject': pdf_meta.get('/Subject', ''),
                        'creator': pdf_meta.get('/Creator', ''),
                        'producer': pdf_meta.get('/Producer', ''),
                        'creation_date': str(pdf_meta.get('/CreationDate', '')),
                        'modification_date': str(pdf_meta.get('/ModDate', ''))
                    })
                
                # 提取每页文本
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            if options.get('preserve_formatting', False):
                                content_parts.append(f"\n--- 第 {page_num + 1} 页 ---\n")
                            content_parts.append(page_text)
                    except Exception as e:
                        content_parts.append(f"\n[第 {page_num + 1} 页提取失败: {str(e)}]\n")
        
        except Exception as e:
            raise Exception(f"PDF处理失败: {str(e)}")
        
        content = "\n".join(content_parts)
        return content, metadata
    
    def _process_text(self, file_path: Path, options: Dict) -> Tuple[str, Dict]:
        """处理纯文本文件"""
        metadata = {
            'format': 'txt',
            'encoding': 'utf-8'
        }
        
        # 尝试不同的编码
        encodings = ['utf-8', 'gbk', 'gb2312', 'gb18030', 'latin-1']
        content = None
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                metadata['encoding'] = encoding
                break
            except UnicodeDecodeError:
                continue
        
        if content is None:
            # 如果所有编码都失败，使用二进制模式读取
            with open(file_path, 'rb') as f:
                content = f.read().decode('utf-8', errors='replace')
            metadata['encoding'] = 'binary_with_replacement'
        
        # 估算页数
        line_count = content.count('\n') + 1
        estimated_pages = max(1, line_count // 50)  # 假设每页50行
        metadata['page_count'] = estimated_pages
        metadata['page_count_note'] = 'estimated'
        metadata['line_count'] = line_count
        
        return content, metadata
    
    def _process_legacy_doc(self, file_path: Path, options: Dict) -> Tuple[str, Dict]:
        """处理旧版DOC文档"""
        # 旧版DOC文档处理比较复杂，MVP版本先作为二进制文件处理
        # 实际项目中可以使用python-docx2txt或其他库
        
        # 尝试作为DOCX处理（有些.doc文件实际上是DOCX格式）
        try:
            return self._process_docx(file_path, options)
        except:
            pass
        
        # 作为文本文件处理
        try:
            return self._process_text(file_path, options)
        except:
            pass
        
        # 如果都失败，返回错误信息
        raise ValueError("无法处理旧版DOC文档，请转换为DOCX格式")
    
    def _analyze_content(self, content: str) -> Dict[str, any]:
        """分析文档内容，提取统计信息"""
        # 基本统计
        word_count = len(content.split())
        char_count = len(content)
        char_count_no_spaces = len(content.replace(' ', '').replace('\n', '').replace('\t', ''))
        
        # 段落统计
        paragraphs = [p for p in content.split('\n\n') if p.strip()]
        paragraph_count = len(paragraphs)
        
        # 句子统计（简单估算）
        sentence_endings = ['.', '!', '?', '。', '！', '？']
        sentence_count = sum(content.count(ending) for ending in sentence_endings)
        
        # 平均值计算
        avg_words_per_paragraph = word_count / paragraph_count if paragraph_count > 0 else 0
        avg_words_per_sentence = word_count / sentence_count if sentence_count > 0 else 0
        
        return {
            'word_count': word_count,
            'character_count': char_count,
            'character_count_no_spaces': char_count_no_spaces,
            'paragraph_count': paragraph_count,
            'sentence_count': sentence_count,
            'avg_words_per_paragraph': round(avg_words_per_paragraph, 1),
            'avg_words_per_sentence': round(avg_words_per_sentence, 1)
        }
    
    def extract_structured_content(self, 
                                 file_path: Union[str, Path],
                                 structure_type: str = 'sections') -> Dict[str, List[str]]:
        """
        提取结构化内容
        
        Args:
            file_path: 文档路径
            structure_type: 结构类型 ('sections', 'headings', 'lists')
        
        Returns:
            Dict: 结构化内容
        """
        result = self.process(file_path)
        if not result.is_success:
            return {}
        
        content = result.content
        structured = {}
        
        if structure_type == 'sections':
            # 按标题分割章节
            # 简单实现：查找可能的标题模式
            sections = []
            current_section = []
            
            for line in content.split('\n'):
                # 检测可能的标题（全大写、数字开头等）
                if (line.isupper() and len(line) > 3) or \
                   (line.strip() and line[0].isdigit() and '.' in line[:3]):
                    if current_section:
                        sections.append('\n'.join(current_section))
                    current_section = [line]
                else:
                    current_section.append(line)
            
            if current_section:
                sections.append('\n'.join(current_section))
            
            structured['sections'] = sections
            
        elif structure_type == 'headings':
            # 提取所有标题
            headings = []
            for line in content.split('\n'):
                line = line.strip()
                if line and (line.isupper() or 
                           (line[0].isdigit() and '.' in line[:3]) or
                           line.startswith('#')):
                    headings.append(line)
            
            structured['headings'] = headings
            
        elif structure_type == 'lists':
            # 提取列表项
            list_items = []
            for line in content.split('\n'):
                line = line.strip()
                if line and (line.startswith(('•', '-', '*', '·')) or
                           (line[0].isdigit() and ('.' in line[:3] or ')' in line[:3]))):
                    list_items.append(line)
            
            structured['lists'] = list_items
        
        return structured
    
    def convert_to_markdown(self, file_path: Union[str, Path]) -> str:
        """
        将文档转换为Markdown格式
        
        Args:
            file_path: 文档路径
        
        Returns:
            str: Markdown格式的内容
        """
        result = self.process(file_path, options={'preserve_formatting': True})
        if not result.is_success:
            return f"# 错误\n\n无法处理文档: {result.error}"
        
        # 添加文档信息头
        markdown_parts = [
            f"# {result.metadata.get('title', '文档内容')}\n",
            f"**文件**: {result.metadata.get('filename', 'unknown')}  ",
            f"**格式**: {result.metadata.get('format', 'unknown')}  ",
            f"**页数**: {result.metadata.get('page_count', 'unknown')}  ",
            f"**字数**: {result.metadata.get('word_count', 'unknown')}  \n",
            "---\n"
        ]
        
        # 处理内容，尝试识别和保留结构
        content = result.content
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                markdown_parts.append('')
                continue
            
            # 识别可能的标题
            if line.isupper() and len(line) > 3:
                markdown_parts.append(f"\n## {line}\n")
            elif line[0:1].isdigit() and '.' in line[:3]:
                # 可能是编号标题
                markdown_parts.append(f"\n### {line}\n")
            else:
                markdown_parts.append(line)
        
        return '\n'.join(markdown_parts)
    
    def get_document_summary(self, file_path: Union[str, Path], max_length: int = 500) -> str:
        """
        获取文档摘要
        
        Args:
            file_path: 文档路径
            max_length: 最大长度
        
        Returns:
            str: 文档摘要
        """
        result = self.process(file_path)
        if not result.is_success:
            return f"无法生成摘要: {result.error}"
        
        content = result.content
        
        # 简单的摘要生成：取前面的内容
        # 在实际应用中，这里可以调用AI模型生成更好的摘要
        if len(content) <= max_length:
            return content
        
        # 尝试在句子边界截断
        truncated = content[:max_length]
        last_period = truncated.rfind('。')
        if last_period == -1:
            last_period = truncated.rfind('.')
        
        if last_period > max_length * 0.8:
            truncated = truncated[:last_period + 1]
        
        return truncated + "..."
    
    def batch_process_documents(self,
                              file_paths: List[Union[str, Path]],
                              options: Optional[Dict] = None) -> List[ProcessingResult]:
        """
        批量处理文档
        
        Args:
            file_paths: 文档路径列表
            options: 处理选项
        
        Returns:
            List[ProcessingResult]: 处理结果列表
        """
        results = []
        options = options or {}
        progress_callback = options.get('progress_callback')
        
        for i, file_path in enumerate(file_paths):
            if progress_callback:
                progress_callback(f"处理文档 {i + 1}/{len(file_paths)}: {Path(file_path).name}")
            
            result = self.process(file_path, options=options)
            results.append(result)
        
        return results
